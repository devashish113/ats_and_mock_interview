"""
Resume Parser Module.
Extracts text content from PDF and DOCX resume files.
"""

import logging
from pathlib import Path
from typing import Optional

# PDF extraction
import fitz  # PyMuPDF

# DOCX extraction
from docx import Document

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_path: Path) -> str:
    """
    Extract text from a PDF file using PyMuPDF.
    
    Args:
        file_path: Path to the PDF file
        
    Returns:
        Extracted text content as a string
    """
    try:
        text_content = []
        
        # Open PDF document
        doc = fitz.open(str(file_path))
        
        # Extract text from each page
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text = page.get_text("text")
            text_content.append(text)
        
        doc.close()
        
        # Join all pages with newlines
        full_text = "\n".join(text_content)
        logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from PDF: {str(e)}")
        raise ValueError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_path: Path) -> str:
    """
    Extract text from a DOCX file using python-docx.
    
    Args:
        file_path: Path to the DOCX file
        
    Returns:
        Extracted text content as a string
    """
    try:
        doc = Document(str(file_path))
        text_content = []
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_content.append(paragraph.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_content.append(" | ".join(row_text))
        
        full_text = "\n".join(text_content)
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"Error extracting text from DOCX: {str(e)}")
        raise ValueError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text(file_path: Path) -> str:
    """
    Extract text from a resume file (PDF or DOCX).
    
    Args:
        file_path: Path to the resume file
        
    Returns:
        Extracted text content as a string
        
    Raises:
        ValueError: If file type is not supported or extraction fails
    """
    extension = file_path.suffix.lower()
    
    if extension == ".pdf":
        return extract_text_from_pdf(file_path)
    elif extension == ".docx":
        return extract_text_from_docx(file_path)
    else:
        raise ValueError(f"Unsupported file format: {extension}. Use PDF or DOCX.")


def clean_text(text: str) -> str:
    """
    Clean extracted text by removing excess whitespace and special characters.
    
    Args:
        text: Raw extracted text
        
    Returns:
        Cleaned text
    """
    import re
    
    # Remove multiple spaces
    text = re.sub(r' +', ' ', text)
    
    # Remove multiple newlines (keep max 2)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # Remove leading/trailing whitespace from each line
    lines = [line.strip() for line in text.split('\n')]
    text = '\n'.join(lines)
    
    return text.strip()

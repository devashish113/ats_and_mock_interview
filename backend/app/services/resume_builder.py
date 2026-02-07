"""
Resume Builder Module.
Generates ATS-friendly DOCX resume files with clean, professional formatting.
Optimized for 1-2 page resumes with human-readable formatting.
"""

import logging
import uuid
import re
from pathlib import Path
from typing import Dict, Optional, List
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE

from app.config import GENERATED_DIR

logger = logging.getLogger(__name__)

# Maximum bullets per section for conciseness
MAX_BULLETS_PER_ROLE = 4

# Strong action verbs for bullet points
ACTION_VERBS = [
    'achieved', 'accomplished', 'administered', 'analyzed', 'architected',
    'automated', 'built', 'collaborated', 'configured', 'created', 'delivered',
    'deployed', 'designed', 'developed', 'drove', 'enhanced', 'established',
    'executed', 'implemented', 'improved', 'increased', 'integrated', 'launched',
    'led', 'managed', 'mentored', 'migrated', 'optimized', 'orchestrated',
    'reduced', 'refactored', 'resolved', 'scaled', 'spearheaded', 'streamlined',
    'transformed', 'upgraded'
]


def clean_text(text: str) -> str:
    """Remove markdown formatting and clean text for DOCX output."""
    if not text:
        return ""
    
    # Remove bold/italic markdown markers
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
    text = re.sub(r'\*(.+?)\*', r'\1', text)
    text = re.sub(r'__(.+?)__', r'\1', text)
    text = re.sub(r'_(.+?)_', r'\1', text)
    
    # Remove markdown headers
    text = re.sub(r'^#{1,6}\s*', '', text, flags=re.MULTILINE)
    
    # Remove markdown links but keep text
    text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
    
    # Remove horizontal rules
    text = re.sub(r'^[-*_]{3,}\s*$', '', text, flags=re.MULTILINE)
    
    # Remove decorative characters
    text = re.sub(r'[★☆●○◆◇■□▪▫►▻→←↑↓]', '', text)
    
    # Clean up multiple spaces and newlines
    text = re.sub(r' +', ' ', text)
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    return text.strip()


def to_title_case(text: str) -> str:
    """Convert text to Title Case."""
    small_words = {'a', 'an', 'and', 'as', 'at', 'but', 'by', 'for', 'in', 
                   'of', 'on', 'or', 'the', 'to', 'with'}
    
    words = text.lower().split()
    result = []
    
    for i, word in enumerate(words):
        if i == 0 or word not in small_words:
            result.append(word.capitalize())
        else:
            result.append(word)
    
    return ' '.join(result)


def add_spacing(paragraph, space_before: int = 0, space_after: int = 6):
    """Add spacing around a paragraph."""
    paragraph.paragraph_format.space_before = Pt(space_before)
    paragraph.paragraph_format.space_after = Pt(space_after)


def is_coursework_line(text: str) -> bool:
    """Check if a line is coursework-related (to filter out)."""
    coursework_keywords = ['coursework', 'relevant courses', 'courses taken', 'course:']
    text_lower = text.lower()
    return any(kw in text_lower for kw in coursework_keywords)


def create_ats_friendly_docx(resume_data: Dict, filename: Optional[str] = None) -> Path:
    """
    Generate an ATS-friendly DOCX resume with refined formatting.
    
    Args:
        resume_data: Dictionary containing structured resume data
        filename: Optional custom filename
        
    Returns:
        Path to the generated DOCX file
    """
    doc = Document()
    
    # Set document margins for 1-2 page fit
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Configure base styles
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    style.paragraph_format.space_after = Pt(0)
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.line_spacing = 1.1
    
    # === HEADER: Name | Phone | Email | GitHub | LinkedIn | Location ===
    _add_header(doc, resume_data)
    
    # === ADD SECTIONS ===
    section_order = ['summary', 'skills', 'experience', 'education', 'projects', 'certifications']
    sections_data = resume_data.get('sections', {})
    
    for section_key in section_order:
        if section_key in sections_data:
            section = sections_data[section_key]
            _add_section(doc, section_key, section)
    
    # Generate filename
    if not filename:
        file_id = str(uuid.uuid4())[:8]
        filename = f"ats_resume_{file_id}.docx"
    
    # Save document
    output_path = GENERATED_DIR / filename
    doc.save(str(output_path))
    
    logger.info(f"Generated ATS resume: {output_path}")
    return output_path


def _add_header(doc: Document, resume_data: Dict):
    """Add refined header with Name | Phone | Email | GitHub | LinkedIn | Location."""
    # === NAME (centered, large, bold) ===
    if resume_data.get('name'):
        name = clean_text(resume_data['name'])
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(name)
        name_run.bold = True
        name_run.font.size = Pt(16)
        name_run.font.name = 'Calibri'
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_spacing(name_para, 0, 3)
    
    # === CONTACT INFO (centered, pipe-separated) ===
    if resume_data.get('contact'):
        contact = resume_data['contact']
        contact_parts = []
        
        # Order: Phone | Email | GitHub | LinkedIn | Location
        if contact.get('phone'):
            contact_parts.append(contact['phone'])
        if contact.get('email'):
            contact_parts.append(contact['email'])
        if contact.get('github'):
            github = contact['github']
            # Clean up URL if full URL
            if 'github.com/' in github:
                github = github.split('github.com/')[-1]
            contact_parts.append(f"github.com/{github}")
        if contact.get('linkedin'):
            linkedin = contact['linkedin']
            if 'linkedin.com/in/' in linkedin:
                linkedin = linkedin.split('linkedin.com/in/')[-1]
            contact_parts.append(f"linkedin.com/in/{linkedin}")
        if contact.get('location'):
            contact_parts.append(contact['location'])
        
        if contact_parts:
            contact_para = doc.add_paragraph()
            contact_run = contact_para.add_run(' | '.join(contact_parts))
            contact_run.font.size = Pt(9.5)
            contact_run.font.name = 'Calibri'
            contact_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_spacing(contact_para, 0, 10)


def _add_section_heading(doc: Document, title: str):
    """Add a bold section heading in Title Case with underline effect."""
    heading_para = doc.add_paragraph()
    heading_run = heading_para.add_run(to_title_case(title))
    heading_run.bold = True
    heading_run.font.size = Pt(11)
    heading_run.font.name = 'Calibri'
    add_spacing(heading_para, 8, 3)
    
    # Add thin separator line
    border_para = doc.add_paragraph()
    border_para.paragraph_format.space_after = Pt(4)


def _add_section(doc: Document, section_key: str, section: Dict):
    """Add a section with refined formatting."""
    section_titles = {
        'summary': 'Professional Summary',
        'skills': 'Skills',
        'experience': 'Work Experience',
        'education': 'Education',
        'projects': 'Projects',
        'certifications': 'Certifications'
    }
    
    title = section_titles.get(section_key, to_title_case(section.get('title', section_key)))
    content = clean_text(section.get('content', ''))
    
    if not content:
        return
    
    # Add section heading
    _add_section_heading(doc, title)
    
    # Format content based on section type
    if section_key == 'skills':
        _add_skills_section(doc, content)
    elif section_key in ['experience', 'projects']:
        _add_experience_section(doc, content, section_key)
    elif section_key == 'education':
        _add_education_section(doc, content)
    elif section_key == 'summary':
        _add_summary_section(doc, content)
    elif section_key == 'certifications':
        _add_certifications_section(doc, content)
    else:
        _add_generic_section(doc, content)


def _add_summary_section(doc: Document, content: str):
    """Add professional summary as a concise paragraph."""
    summary_text = ' '.join(line.strip() for line in content.split('\n') if line.strip())
    summary_text = clean_text(summary_text)
    
    # Limit summary length for 1-2 page fit
    if len(summary_text) > 500:
        summary_text = summary_text[:497] + '...'
    
    para = doc.add_paragraph()
    run = para.add_run(summary_text)
    run.font.size = Pt(10.5)
    run.font.name = 'Calibri'
    add_spacing(para, 0, 6)


def _add_skills_section(doc: Document, content: str):
    """
    Add skills in grouped format:
    Technical Skills: Python, Java, SQL, React, etc.
    Soft Skills: Leadership, Communication, etc.
    """
    lines = content.split('\n')
    
    technical_skills = []
    soft_skills = []
    other_skills = []
    
    soft_skill_keywords = ['communication', 'leadership', 'teamwork', 'problem-solving',
                          'collaboration', 'management', 'analytical', 'critical thinking',
                          'time management', 'adaptability', 'creativity', 'interpersonal']
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        # Remove bullet characters
        line = re.sub(r'^[-•*→●○]\s*', '', line)
        line = clean_text(line)
        
        if not line:
            continue
        
        # Check if line has category prefix (e.g., "Technical: Python, Java")
        if ':' in line:
            parts = line.split(':', 1)
            category = parts[0].lower()
            skills = parts[1].strip()
            
            if 'technical' in category or 'programming' in category or 'tool' in category:
                technical_skills.append(skills)
            elif 'soft' in category or 'interpersonal' in category:
                soft_skills.append(skills)
            else:
                other_skills.append(line)
        else:
            # Categorize based on content
            line_lower = line.lower()
            if any(kw in line_lower for kw in soft_skill_keywords):
                soft_skills.append(line)
            else:
                technical_skills.append(line)
    
    # Output grouped skills
    if technical_skills:
        tech_text = ', '.join(technical_skills)
        # Clean up duplicate commas
        tech_text = re.sub(r',\s*,', ',', tech_text)
        tech_text = re.sub(r'\s+', ' ', tech_text)
        
        para = doc.add_paragraph()
        label = para.add_run('Technical Skills: ')
        label.bold = True
        label.font.size = Pt(10.5)
        label.font.name = 'Calibri'
        content_run = para.add_run(tech_text)
        content_run.font.size = Pt(10.5)
        content_run.font.name = 'Calibri'
        add_spacing(para, 0, 3)
    
    if soft_skills:
        soft_text = ', '.join(soft_skills)
        soft_text = re.sub(r',\s*,', ',', soft_text)
        soft_text = re.sub(r'\s+', ' ', soft_text)
        
        para = doc.add_paragraph()
        label = para.add_run('Soft Skills: ')
        label.bold = True
        label.font.size = Pt(10.5)
        label.font.name = 'Calibri'
        content_run = para.add_run(soft_text)
        content_run.font.size = Pt(10.5)
        content_run.font.name = 'Calibri'
        add_spacing(para, 0, 3)
    
    if other_skills:
        for skill_line in other_skills:
            para = doc.add_paragraph()
            run = para.add_run(skill_line)
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
            add_spacing(para, 0, 2)


def _add_experience_section(doc: Document, content: str, section_type: str):
    """Add experience/projects with max 3-4 bullets per role."""
    lines = content.split('\n')
    current_role = None
    bullet_count = 0
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        line = clean_text(line)
        is_bullet = bool(re.match(r'^[-•*→●○]\s*', line))
        
        if is_bullet:
            # Only add up to MAX_BULLETS_PER_ROLE bullets per role
            if bullet_count >= MAX_BULLETS_PER_ROLE:
                continue
            
            bullet_text = re.sub(r'^[-•*→●○]\s*', '', line).strip()
            if bullet_text:
                para = doc.add_paragraph()
                para.paragraph_format.left_indent = Inches(0.2)
                run = para.add_run(f"• {bullet_text}")
                run.font.size = Pt(10.5)
                run.font.name = 'Calibri'
                add_spacing(para, 0, 2)
                bullet_count += 1
        else:
            # This is a role/company header or project name
            # Reset bullet count for new role
            bullet_count = 0
            
            para = doc.add_paragraph()
            # Check if it looks like a job entry (contains | or dates)
            if '|' in line or re.search(r'\d{4}', line):
                run = para.add_run(line)
                run.bold = True
            else:
                run = para.add_run(line)
                if section_type == 'projects':
                    run.bold = True
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
            add_spacing(para, 3, 2)


def _add_education_section(doc: Document, content: str):
    """Add education entries as clean compact lines."""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Remove bullet characters
        line = re.sub(r'^[-•*→●○]\s*', '', line)
        line = clean_text(line)
        
        # Skip coursework lines
        if is_coursework_line(line):
            continue
        
        if line:
            para = doc.add_paragraph()
            # Bold the degree line
            if any(deg in line.lower() for deg in ['bachelor', 'master', 'phd', 'b.tech', 'm.tech', 'b.s', 'm.s', 'bba', 'mba']):
                run = para.add_run(line)
                run.bold = True
            else:
                run = para.add_run(line)
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
            add_spacing(para, 0, 2)


def _add_certifications_section(doc: Document, content: str):
    """Add certifications, excluding coursework lines."""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # Remove bullet characters
        line = re.sub(r'^[-•*→●○]\s*', '', line)
        line = clean_text(line)
        
        # Skip coursework lines
        if is_coursework_line(line):
            continue
        
        if line:
            para = doc.add_paragraph()
            run = para.add_run(f"• {line}")
            run.font.size = Pt(10.5)
            run.font.name = 'Calibri'
            add_spacing(para, 0, 2)


def _add_generic_section(doc: Document, content: str):
    """Add generic section content."""
    lines = content.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        line = clean_text(line)
        is_bullet = bool(re.match(r'^[-•*→●○]\s*', line))
        
        if is_bullet:
            bullet_text = re.sub(r'^[-•*→●○]\s*', '', line).strip()
            para = doc.add_paragraph()
            para.paragraph_format.left_indent = Inches(0.2)
            run = para.add_run(f"• {bullet_text}")
        else:
            para = doc.add_paragraph()
            run = para.add_run(line)
        
        run.font.size = Pt(10.5)
        run.font.name = 'Calibri'
        add_spacing(para, 0, 2)


def generate_from_ai_rewrite(rewritten_text: str, original_data: Dict, filename: Optional[str] = None) -> Path:
    """
    Generate DOCX from AI-rewritten resume text with refined formatting.
    
    Args:
        rewritten_text: AI-rewritten resume text
        original_data: Original resume data (for name/contact)
        filename: Optional custom filename
        
    Returns:
        Path to generated DOCX file
    """
    doc = Document()
    
    # Set document margins for compact fit
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.6)
        section.right_margin = Inches(0.6)
    
    # Configure base style
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10.5)
    
    # Clean the rewritten text
    rewritten_text = clean_text(rewritten_text)
    
    # === HEADER ===
    _add_header(doc, original_data)
    
    # Parse rewritten content into sections
    lines = rewritten_text.split('\n')
    
    section_keywords = {
        'summary': ['summary', 'objective', 'profile', 'about'],
        'skills': ['skills', 'competencies', 'expertise', 'proficiencies'],
        'experience': ['experience', 'employment', 'work history', 'internship'],
        'education': ['education', 'academic', 'qualifications'],
        'projects': ['projects', 'portfolio'],
        'certifications': ['certifications', 'licenses', 'training']
    }
    
    current_section = None
    section_content = []
    name_added = bool(original_data.get('name'))
    contact_added = bool(original_data.get('contact'))
    
    for line in lines:
        original_line = line
        line = line.strip()
        
        if not line:
            continue
        
        # Skip duplicate name/contact
        if name_added and original_data.get('name') and line.lower() == original_data['name'].lower():
            continue
        if contact_added and '@' in line and len(line.split()) <= 5:
            continue
        
        # Check if section header
        is_section_header = False
        detected_section = None
        
        for sec_key, keywords in section_keywords.items():
            for keyword in keywords:
                if keyword in line.lower() and len(line) < 50:
                    is_section_header = True
                    detected_section = sec_key
                    break
            if is_section_header:
                break
        
        if is_section_header and detected_section:
            if current_section and section_content:
                _add_parsed_section(doc, current_section, section_content)
            current_section = detected_section
            section_content = []
        else:
            section_content.append(original_line)
    
    # Add final section
    if current_section and section_content:
        _add_parsed_section(doc, current_section, section_content)
    
    # Generate filename
    if not filename:
        file_id = str(uuid.uuid4())[:8]
        filename = f"ats_resume_{file_id}.docx"
    
    output_path = GENERATED_DIR / filename
    doc.save(str(output_path))
    
    logger.info(f"Generated ATS resume from AI rewrite: {output_path}")
    return output_path


def _add_parsed_section(doc: Document, section_key: str, content_lines: List[str]):
    """Add a parsed section with refined formatting."""
    section_titles = {
        'summary': 'Professional Summary',
        'skills': 'Skills',
        'experience': 'Work Experience',
        'education': 'Education',
        'projects': 'Projects',
        'certifications': 'Certifications'
    }
    
    title = section_titles.get(section_key, to_title_case(section_key))
    _add_section_heading(doc, title)
    
    if section_key == 'skills':
        # Group skills into Technical and Soft
        technical_skills = []
        soft_skills = []
        
        soft_keywords = ['communication', 'leadership', 'teamwork', 'problem-solving',
                        'collaboration', 'management', 'analytical', 'critical thinking']
        
        for line in content_lines:
            line = line.strip()
            if line:
                line = re.sub(r'^[-•*→●○]\s*', '', line)
                line = clean_text(line)
                if line:
                    if any(kw in line.lower() for kw in soft_keywords):
                        soft_skills.append(line)
                    else:
                        technical_skills.append(line)
        
        if technical_skills:
            para = doc.add_paragraph()
            label = para.add_run('Technical Skills: ')
            label.bold = True
            label.font.size = Pt(10.5)
            content_run = para.add_run(', '.join(technical_skills))
            content_run.font.size = Pt(10.5)
            add_spacing(para, 0, 3)
        
        if soft_skills:
            para = doc.add_paragraph()
            label = para.add_run('Soft Skills: ')
            label.bold = True
            label.font.size = Pt(10.5)
            content_run = para.add_run(', '.join(soft_skills))
            content_run.font.size = Pt(10.5)
            add_spacing(para, 0, 3)
    
    elif section_key == 'summary':
        summary_text = ' '.join(line.strip() for line in content_lines if line.strip())
        summary_text = clean_text(summary_text)
        if len(summary_text) > 500:
            summary_text = summary_text[:497] + '...'
        if summary_text:
            para = doc.add_paragraph()
            run = para.add_run(summary_text)
            run.font.size = Pt(10.5)
            add_spacing(para, 0, 6)
    
    elif section_key in ['experience', 'projects']:
        bullet_count = 0
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
            
            line = clean_text(line)
            is_bullet = bool(re.match(r'^[-•*→●○]\s*', line))
            
            if is_bullet:
                if bullet_count >= MAX_BULLETS_PER_ROLE:
                    continue
                bullet_text = re.sub(r'^[-•*→●○]\s*', '', line).strip()
                if bullet_text:
                    para = doc.add_paragraph()
                    para.paragraph_format.left_indent = Inches(0.2)
                    run = para.add_run(f"• {bullet_text}")
                    run.font.size = Pt(10.5)
                    add_spacing(para, 0, 2)
                    bullet_count += 1
            else:
                bullet_count = 0  # Reset for new role
                para = doc.add_paragraph()
                if '|' in line or re.search(r'\d{4}', line):
                    run = para.add_run(line)
                    run.bold = True
                else:
                    run = para.add_run(line)
                    run.bold = True
                run.font.size = Pt(10.5)
                add_spacing(para, 3, 2)
    
    elif section_key == 'education':
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
            line = re.sub(r'^[-•*→●○]\s*', '', line)
            line = clean_text(line)
            if is_coursework_line(line):
                continue
            if line:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(10.5)
                add_spacing(para, 0, 2)
    
    elif section_key == 'certifications':
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
            line = re.sub(r'^[-•*→●○]\s*', '', line)
            line = clean_text(line)
            if is_coursework_line(line):
                continue
            if line:
                para = doc.add_paragraph()
                run = para.add_run(f"• {line}")
                run.font.size = Pt(10.5)
                add_spacing(para, 0, 2)
    
    else:
        for line in content_lines:
            line = line.strip()
            if not line:
                continue
            line = re.sub(r'^[-•*→●○]\s*', '', line)
            line = clean_text(line)
            if line:
                para = doc.add_paragraph()
                run = para.add_run(line)
                run.font.size = Pt(10.5)
                add_spacing(para, 0, 2)


def get_download_path(file_id: str) -> Optional[Path]:
    """Get the path to a generated resume file."""
    file_path = GENERATED_DIR / f"{file_id}.docx"
    if file_path.exists():
        return file_path
    
    file_path = GENERATED_DIR / file_id
    if file_path.exists():
        return file_path
    
    for f in GENERATED_DIR.glob(f"*{file_id}*"):
        if f.is_file():
            return f
    
    return None
